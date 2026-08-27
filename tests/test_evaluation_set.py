from __future__ import annotations

import tempfile
import unittest
from pathlib import Path
from types import SimpleNamespace

from docx import Document
from docx.shared import RGBColor
from candidate_profile import CandidateProfileData, EligibilityAssessment, EligibilityCheck
from models.evaluation import WEIGHTS

from services.evaluation_set import (
    EvaluationCase,
    _all_gap_text,
    aggregate_case_runs,
    failed_case_result,
    gate_expectation_passes,
    import_cases_from_docx,
    run_evaluation_case,
)


class EvaluationSetTests(unittest.TestCase):
    def test_expected_gap_can_be_recalled_from_professional_gate(self):
        evaluation = SimpleNamespace(
            missing_keywords=[],
            risks=[],
            **{
                field: SimpleNamespace(gaps=[])
                for field in WEIGHTS
            },
        )
        eligibility = EligibilityAssessment(
            checks=[
                EligibilityCheck(
                    constraint_type="professional_skill",
                    requirement="熟练使用 SQL",
                    result="不满足",
                    candidate_evidence="仅掌握 SQL 基础查询",
                )
            ]
        )

        self.assertIn("sql", _all_gap_text(evaluation, eligibility))

    def test_incomplete_case_is_expected_to_abstain(self):
        case = EvaluationCase(
            case_id="case_incomplete",
            title="示例公司｜商品运营实习生",
            company="示例公司",
            role="商品运营实习生",
            jd_text=(
                "这张截图只展示了岗位要求，没有看到完整的岗位职责。\n"
                "任职要求：本科及以上学历，具备数据分析能力。"
            ),
            human_judgment="信息不完整，不应评分",
            score_min=0,
            score_max=0,
            expected_gate="非不满足",
            expected_gap_groups=[],
            expected_scorable=False,
        )
        self.assertFalse(case.expected_scorable)
        result = run_evaluation_case(
            case,
            client=None,
            model="unused",
            resume_text="unused",
            profile=CandidateProfileData(),
        )
        self.assertTrue(result["abstained"])
        self.assertTrue(result["quality_pass"])

    def test_imports_ten_red_judgments_from_docx(self):
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "cases.docx"
            destination = Path(temp_dir) / "cases.json"
            document = Document()
            for number in range(1, 11):
                document.add_heading(f"{number}. 示例公司｜岗位{number}", level=1)
                document.add_heading("工作职责", level=3)
                document.add_paragraph("岗位职责：负责完整的行业研究和数据分析工作。")
                document.add_heading("任职要求", level=3)
                document.add_paragraph("任职要求：本科及以上学历，具备结构化思维。")
                judgment = document.add_paragraph()
                run = judgment.add_run("判断：70-84，方向基本匹配")
                run.font.color.rgb = RGBColor(192, 0, 0)
            document.save(source)

            cases = import_cases_from_docx(source, destination)

            self.assertEqual(len(cases), 10)
            self.assertEqual(cases[0].score_min, 70)
            self.assertEqual(cases[0].score_max, 84)
            self.assertTrue(destination.exists())

    def test_aggregate_requires_range_gate_gap_and_stability(self):
        case = EvaluationCase(
            case_id="case_001",
            title="示例",
            company="示例",
            role="战略实习生",
            jd_text="完整 JD",
            human_judgment="方向匹配",
            score_min=70,
            score_max=84,
            expected_gate="非不满足",
            expected_gap_groups=[],
        )
        runs = [
            {
                "actual_score": score,
                "actual_gate": "满足",
                "gate_pass": True,
                "gap_pass": True,
                "score_pass": True,
                "case_id": "case_001",
                "title": "示例",
                "expected_score": "70-84",
            }
            for score in (76, 78, 80)
        ]
        result = aggregate_case_runs(case, runs)
        self.assertEqual(result["actual_score"], 78)
        self.assertEqual(result["score_spread"], 4)
        self.assertEqual(result["run_scores"], [76, 78, 80])
        self.assertTrue(result["passed"])

    def test_failed_case_has_complete_report_shape(self):
        case = EvaluationCase(
            case_id="case_003",
            title="示例",
            company="示例",
            role="运营实习生",
            jd_text="完整 JD",
            human_judgment="部分匹配",
            score_min=60,
            score_max=75,
            expected_gate="非不满足",
            expected_gap_groups=[],
        )
        result = failed_case_result(case, error="模型格式错误", model="test")
        self.assertFalse(result["passed"])
        self.assertIsNone(result["actual_score"])
        self.assertEqual(result["actual_gate"], "未完成")
        self.assertIn("模型格式错误", result["error"])

    def test_exact_satisfied_gate_does_not_accept_unknown(self):
        self.assertTrue(gate_expectation_passes("满足", "满足"))
        self.assertFalse(gate_expectation_passes("满足", "存疑"))
        self.assertTrue(gate_expectation_passes("非不满足", "存疑"))


if __name__ == "__main__":
    unittest.main()
